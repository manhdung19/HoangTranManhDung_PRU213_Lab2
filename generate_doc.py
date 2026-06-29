import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level, space_before=12, space_after=6):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(space_before)
    heading.paragraph_format.space_after = Pt(space_after)
    heading.paragraph_format.keep_with_next = True
    
    # Style runs
    for run in heading.runs:
        run.font.name = 'Arial'
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(26, 82, 118) # Deep Blue
            run.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(41, 128, 185) # Medium Blue
            run.bold = True
        else:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(52, 73, 94) # Dark Slate
            run.bold = True
    return heading

def create_document():
    doc = Document()
    
    # Page setup - Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base Font Settings
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(44, 62, 80) # Dark Charcoal
    
    # ------------------- TITLE PAGE -------------------
    # Add spacing at the top
    for _ in range(3):
        doc.add_paragraph()
        
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("FPT UNIVERSITY\nDEPARTMENT OF COMPUTER SCIENCE\n\n")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = RGBColor(127, 140, 141)
    title_run.bold = True
    
    title_p2 = doc.add_paragraph()
    title_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p2.paragraph_format.space_after = Pt(24)
    t2_run = title_p2.add_run("PRU213 - GAME DEVELOPMENT WITH C# AND UNITY\n")
    t2_run.font.name = 'Arial'
    t2_run.font.size = Pt(18)
    t2_run.bold = True
    t2_run.font.color.rgb = RGBColor(26, 82, 118)
    
    t3_run = title_p2.add_run("LAB 2 ASSIGNMENT: 2D SNOWBOARDER GAME\n")
    t3_run.font.name = 'Arial'
    t3_run.font.size = Pt(22)
    t3_run.bold = True
    t3_run.font.color.rgb = RGBColor(231, 76, 60) # Coral/Orange-red
    
    for _ in range(4):
        doc.add_paragraph()
        
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_p.add_run(
        "Student Name: Hoang Tran Manh Dung\n"
        "Student ID: [Your Student ID]\n"
        "Class: PRU213\n"
        "Project Title: Snow Boarder 2D\n"
        "Instructor: FPT Academic Board\n"
        "Date of Submission: June 2026\n"
    )
    info_run.font.name = 'Calibri'
    info_run.font.size = Pt(12)
    info_run.bold = True
    
    doc.add_page_break()
    
    # ------------------- SECTION 1 -------------------
    add_heading_styled(doc, "1. Project Overview & Objective", level=1)
    
    p = doc.add_paragraph(
        "The 'Snow Boarder 2D' project is an action-adventure sports game built using Unity 6 and C#. "
        "The game allows the player to control a character named Barry as he snowboards down various slopes, "
        "navigates steep hills, avoids physical obstacles, performs mid-air rotation tricks, and collects "
        "points and power-ups to maximize their high score. The main goal of this Lab 2 assignment is to apply "
        "advanced 2D physics engine features in Unity, design responsive user controls, build modular C# scripts, "
        "and establish a functional HUD and game state management system."
    )
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.15

    # ------------------- SECTION 2 -------------------
    add_heading_styled(doc, "2. Game Features and Mechanics", level=1)
    
    add_heading_styled(doc, "2.1 Control System & User Inputs", level=2)
    p = doc.add_paragraph(
        "The control system is designed to be responsive, tactile, and intuitive. "
        "Players can control the snowboarder's balance, rotation, and speed using the keyboard:"
    )
    p.paragraph_format.space_after = Pt(6)
    
    controls = [
        ("Left Arrow / A Key", "Rotates the player counter-clockwise (performs Backflips in mid-air)."),
        ("Right Arrow / D Key", "Rotates the player clockwise (performs Frontflips in mid-air)."),
        ("Up Arrow / Spacebar", "Accelerates the snowboarder down the slope (Speed Boost)."),
        ("Down Arrow / S Key", "Applies friction and slows down the snowboarder (Brakes).")
    ]
    for key, desc in controls:
        bullet = doc.add_paragraph(style='List Bullet')
        r1 = bullet.add_run(key + ": ")
        r1.bold = True
        bullet.add_run(desc)
        bullet.paragraph_format.space_after = Pt(4)
        
    add_heading_styled(doc, "2.2 Physics-Based Slopes and Obstacles", level=2)
    p = doc.add_paragraph(
        "Slopes are created using Unity's 2D Sprite Shape package with custom Bezier curves, producing a natural, "
        "smooth terrain. The movement of the player is driven by a SurfaceEffector2D attached to the terrain, "
        "providing a constant forwarding force along the slope. To prevent friction lag, a low-friction "
        "PhysicsMaterial2D is applied to the player's snowboard.\n\n"
        "Physical obstacles (such as rocks) are placed along the path. These obstacles utilize a PolygonCollider2D "
        "shaped to match the rock's outline and a zero-friction PhysicsMaterial2D. Additionally, a local "
        "SurfaceEffector2D is added to the rock's collider, allowing the player to naturally slide over the top of the "
        "rock rather than getting physically stuck, while simultaneously suffering a speed penalty as punishment."
    )
    p.paragraph_format.space_after = Pt(12)

    add_heading_styled(doc, "2.3 Mid-Air Trick & Rotation System", level=2)
    p = doc.add_paragraph(
        "When the player launches off a ramp or cliff, the C# script monitors the cumulative rotation angle (Z-axis). "
        "If the player successfully completes a full 360-degree rotation (or multiple of 360) and lands on their board "
        "without crashing, they receive a massive score bonus and their score multiplier (combo) increases. "
        "If the player lands on their head, they crash, lose a life, and the combo multiplier resets to x1."
    )
    p.paragraph_format.space_after = Pt(12)

    add_heading_styled(doc, "2.4 Collectibles & Power-up States", level=2)
    p = doc.add_paragraph(
        "There are three types of interactive items placed throughout the level, utilizing Trigger Colliders to avoid blocking movement:"
    )
    p.paragraph_format.space_after = Pt(6)
    
    items = [
        ("Snowflakes (Blue Sprite)", "Adds direct points to the player's score. The score awarded is multiplied by the current combo factor (e.g., Score + 10 * Combo)."),
        ("Speed Boost (Yellow Sprite)", "Boosts the terrain's SurfaceEffector2D speed to 1.5x of the normal boost speed for a duration of 3 seconds, enabling massive jumps."),
        ("Invincibility (Gold Sprite)", "Renders Barry completely immune to crashes (head-ground contact) and obstacle speed reductions for 5 seconds. During this period, Barry's sprite flashes bright yellow to indicate invulnerability.")
    ]
    for item, desc in items:
        bullet = doc.add_paragraph(style='List Bullet')
        r1 = bullet.add_run(item + ": ")
        r1.bold = True
        bullet.add_run(desc)
        bullet.paragraph_format.space_after = Pt(4)

    add_heading_styled(doc, "2.5 Lives & Respawning Checkpoint System", level=2)
    p = doc.add_paragraph(
        "The player begins the game with 3 lives. If the player crashes (heads hits the ground), they lose 1 life. "
        "To prevent repetitive frustration, the game features checkpoint flags. When a player passes a checkpoint, "
        "their respawn location updates to that checkpoint flag. Upon crashing, the player's Rigidbody2D is reset, "
        "and the player respawns 1.5 meters above the checkpoint to prevent clipping or getting stuck in the ground. "
        "If lives reach 0, the game halts, and the Game Over screen is displayed."
    )
    p.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # ------------------- SECTION 3 -------------------
    add_heading_styled(doc, "3. Software Architecture and Scripts", level=1)
    
    p = doc.add_paragraph(
        "The software architecture is modular and decoupled, consisting of specialized scripts managing physics, "
        "UI, game logic, and collisions. Below is a detailed breakdown of each script in the C# codebase:"
    )
    p.paragraph_format.space_after = Pt(12)

    # Create a table for the scripts
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Shading Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Script Name'
    hdr_cells[1].text = 'Primary Role / Functionality'
    hdr_cells[2].text = 'Key Methods & Events'
    
    for cell in hdr_cells:
        set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            
    script_data = [
        ("PlayerController.cs", 
         "Main player script. Handles movement physics (adding torque, controlling SurfaceEffector2D speed based on input), checking if grounded, air trick rotation angles, score multipliers, lives, checkpoints, and power-up coroutines.",
         "RespondToBoost(), RotatePlayer(), CalculateSpeedScore(), TrackAirRotation(), AddCollectibleScore(), ActivateSpeedBoost(), ActivateInvincibility(), CrashPlayer(), Respawn()"),
         
        ("CrashDetector.cs",
         "Attached to Barry's head collider. Triggers a crash event if the head collides with the ground, triggering audio and invoking the player's CrashPlayer() method.",
         "OnCollisionEnter2D()"),
         
        ("Collectible.cs",
         "Attached to Snowflake and Power-up prefabs. Detects trigger overlap by the player, calls the corresponding benefit method on the PlayerController, and destroys the collectible.",
         "OnTriggerEnter2D(), Destroy(gameObject)"),
         
        ("Checkpoint.cs",
         "Updates the player's respawn position to the checkpoint's coordinates when the player triggers the flag, playing a visual/audio indicator.",
         "OnTriggerEnter2D()"),
         
        ("FinishLine.cs",
         "Triggers level completion when Barry crosses the finish line. Halts controls, plays victory particles/audio, and opens the victory UI.",
         "OnTriggerEnter2D()"),
         
        ("HUDManager.cs",
         "Refreshes the Head-Up Display (UI) elements on every frame, showing real-time Score, current Combo multiplier, Speed in km/h, and Lives.",
         "Update(), RefreshHUD()"),
         
        ("MainMenuController.cs",
         "Handles Main Menu interactions, loading Level 1 scene, displaying options, and quitting the application.",
         "PlayGame(), QuitGame()")
    ]
    
    for s_name, s_role, s_methods in script_data:
        row_cells = table.add_row().cells
        row_cells[0].text = s_name
        row_cells[1].text = s_role
        row_cells[2].text = s_methods
        
        for cell in row_cells:
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            cell.paragraphs[0].runs[0].font.name = 'Calibri'
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            
        # Bold script name
        row_cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ------------------- SECTION 4 -------------------
    add_heading_styled(doc, "4. Implementation & Design Choices", level=1)
    
    add_heading_styled(doc, "4.1 Visual Styling & UI", level=2)
    p = doc.add_paragraph(
        "A premium look is achieved using a cohesive color palette: white slopes, soft blue fillings, and custom "
        "backgrounds. A high-quality 2D cartoon background ('menu_bg.png') has been integrated into the Main Menu Canvas, "
        "paired with TextMeshPro components using modern typography to ensure a professional finish. "
        "The HUD display uses a text-based hearts indicator for lives (e.g. 'Lives: 3') to prevent unsupported character "
        "box renders in default TMPro assets."
    )
    p.paragraph_format.space_after = Pt(12)

    add_heading_styled(doc, "4.2 Code Optimization & Bug Fixes", level=2)
    p = doc.add_paragraph(
        "During development, several key bugs were identified and optimized:\n\n"
        "1. Input Lag: Character rotation felt sluggish. This was resolved by reducing the Rigidbody2D's "
        "Angular Damping (Angular Drag in Unity 6) from 15 to 0.05, allowing responsive and snappy rotation.\n"
        "2. Physics Tunneling: Barry passed through snowflakes without eating them. This was fixed by increasing the "
        "CircleCollider2D Radius of the snowflake prefab to 8.0, compensating for the small scale (0.06) and preventing "
        "high-speed frame skips.\n"
        "3. Obstacle Friction: Barry got stuck on top of rock obstacles. This was fixed by adding a zero-friction "
        "PhysicsMaterial2D to the rock and applying a SurfaceEffector2D to the rock collider, allowing the player "
        "to slide smoothly over the rock surface while applying the intended speed penalty.\n"
        "4. Multiple Effectors Bug: Finding the SurfaceEffector2D dynamically using FindFirstObjectByType caused the script "
        "to randomly bind to the rock's effector instead of the slope's effector. This was resolved by explicitly assigning "
        "the main Level Sprite Shape's SurfaceEffector2D in the PlayerController component's inspector."
    )
    p.paragraph_format.space_after = Pt(12)

    doc.save("Lab2_Project_Documentation.docx")
    print("Document successfully created!")

if __name__ == "__main__":
    create_document()
